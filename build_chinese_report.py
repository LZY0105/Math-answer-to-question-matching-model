from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "17365D"
TEAL = "2F6B73"
GOLD = "B08D57"
INK = "20252B"
MUTED = "5F6B76"
LIGHT = "F2F5F7"
LIGHT_BLUE = "E8EEF5"
WHITE = "FFFFFF"
RED = "9B1C1C"

BODY_CJK = "宋体"
BODY_LATIN = "Times New Roman"
HEAD_CJK = "微软雅黑"
MONO = "Consolas"


TABLE_TITLES = [
    "相关对齐系统比较",
    "四阶段级联的信号、适用条件与成本",
    "文本层质量状态及处理方式",
    "置信度等级与证据基础",
    "评估语料构成",
    "保留完整文档结构时的端到端结果",
    "书签消融实验结果",
    "书签消融条件下的页面级延迟",
    "运算符上下文窗口半径扫描",
    "位置先验在不同文档条件下的表现",
    "复现实验所含测试套件",
    "系统实测常量及其来源",
]
FIGURE_RE = re.compile(
    r"(?:\[\[FIGURE_([1-4])\]\]|!\[[^\]]*\]\(figures/cn_report/figure_([1-4])\.png\))"
)


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="AAB4BE", size="5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_run_font(run, cjk=BODY_CJK, latin=BODY_LATIN, size=10.5, bold=None, italic=None, color=INK):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), latin)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color=NAVY, size="14", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill=LIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, cjk=HEAD_CJK, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_LATIN)
    r_fonts.set(qn("w:hAnsi"), BODY_LATIN)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")
    r_pr.append(sz)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(r_pr)
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, cjk=HEAD_CJK, size=8.5, color=MUTED)


def add_numbering_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract or [0]) + 1
    num_id = max(existing_num or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(ind)
    for node in (start, num_fmt, lvl_text, suff, p_pr):
        lvl.append(node)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_inline(paragraph, text, size=10.5, color=INK, cjk=BODY_CJK, latin=BODY_LATIN):
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, cjk=cjk, latin=latin, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, cjk=cjk, latin=latin, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, cjk=HEAD_CJK, latin=MONO, size=size - 0.3, color=TEAL)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "EEF2F3")
            run._element.get_or_add_rPr().append(shade)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, cjk=cjk, latin=latin, size=size, color=color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, cjk=cjk, latin=latin, size=size, color=color)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK)
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74)
    pf.widow_control = True

    for style_name, size, before, after, color in (
        ("Heading 1", 15, 16, 8, NAVY),
        ("Heading 2", 12.5, 12, 6, TEAL),
        ("Heading 3", 11.5, 9, 4, NAVY),
    ):
        style = styles[style_name]
        style.font.name = BODY_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    caption = styles["Caption"]
    caption.font.name = BODY_LATIN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK)
    caption.font.size = Pt(9)
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.keep_together = True

    if "Abstract Body" not in styles:
        style = styles.add_style("Abstract Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Abstract Body"]
    style.font.name = BODY_LATIN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK)
    style.font.size = Pt(10.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.45
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Cm(0.74)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.55)
    section.bottom_margin = Cm(2.45)
    section.left_margin = Cm(2.75)
    section.right_margin = Cm(2.45)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("FIND-ENGINE　技术研究报告")
    set_run_font(run, cjk=HEAD_CJK, size=8.5, bold=True, color=MUTED)
    set_paragraph_border_bottom(p, color="CBD3DA", size="4", space="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    add_page_field(p)



def read_version(markdown_path: Path):
    """Version and date come from the report, never from this file.

    They were hardcoded here as 1.0 / 2026-08-25 and stayed there while the
    report moved on, so the built paper contradicted its own source. The
    markdown line looks like:  *版本 2.0｜2026 年 8 月 28 日*
    """
    for line in markdown_path.read_text(encoding="utf-8").splitlines()[:20]:
        m = re.match(r"^\*\s*(版本[^｜|]*)[｜|]\s*(.+?)\s*\*$", line.strip())
        if m:
            return m.group(1).strip(), m.group(2).strip()
    raise SystemExit("no '*版本 X｜日期*' line found in %s" % markdown_path)

def add_cover(doc: Document, version: str, date: str):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("技术研究报告")
    set_run_font(run, cjk=HEAD_CJK, size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    run = p.add_run("Find-Engine")
    set_run_font(run, cjk=HEAD_CJK, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_together = True
    run = p.add_run("面向成对数学 PDF 的\n确定性题目-答案对齐")
    set_run_font(run, cjk=HEAD_CJK, size=18, bold=True, color=TEAL)
    set_paragraph_border_bottom(p, color=GOLD, size="12", space="12")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(version)
    set_run_font(run, cjk=HEAD_CJK, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(date)
    set_run_font(run, cjk=HEAD_CJK, size=10.5, color=MUTED)

    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(13)
    p.add_run().add_break(WD_BREAK.PAGE)


def choose_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 2:
        max_lens = [max(len(r[i]) for r in rows) for i in range(cols)]
        if max_lens[0] < 16:
            return [2400, 6960]
        return [4300, 5060]
    if cols == 3:
        return [2400, 1800, 5160]
    if cols == 4:
        return [1750, 2550, 2750, 2310]
    if cols == 7:
        return [1800, 1600, 1600, 900, 800, 1100, 1560]
    total = 9360
    base = total // cols
    return [base] * (cols - 1) + [total - base * (cols - 1)]


def add_table(doc: Document, rows: list[list[str]], number: int):
    title = TABLE_TITLES[number - 1] if number <= len(TABLE_TITLES) else f"数据汇总 {number}"
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(5)
    cap.paragraph_format.keep_with_next = True
    run = cap.add_run(f"表 {number}　{title}")
    set_run_font(run, cjk=HEAD_CJK, size=9, bold=True, color=NAVY)

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = choose_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            numeric = bool(re.fullmatch(r"[\d\s.,/%*†—-]+", re.sub(r"\*", "", value)))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if numeric or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(
                p,
                value,
                size=8.6 if cols >= 4 else 9.0,
                color=WHITE if r_idx == 0 else INK,
                cjk=HEAD_CJK if r_idx == 0 else BODY_CJK,
            )
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.line_spacing = 1


def make_figures(out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\simsun.ttc")

    def font(size, bold=False):
        bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
        selected = bold_path if bold and bold_path.exists() else font_path
        return ImageFont.truetype(str(selected), size=size)

    def centered(draw, xy, text, ft, fill, anchor="mm"):
        draw.text(xy, text, font=ft, fill=fill, anchor=anchor)

    def arrow(draw, start, end, color="#7B8792", width=4):
        draw.line([start, end], fill=color, width=width)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            sign = 1 if ex > sx else -1
            pts = [(ex, ey), (ex - sign * 18, ey - 10), (ex - sign * 18, ey + 10)]
        else:
            sign = 1 if ey > sy else -1
            pts = [(ex, ey), (ex - 10, ey - sign * 18), (ex + 10, ey - sign * 18)]
        draw.polygon(pts, fill=color)

    # Figure 1: four-stage cascade.
    image = Image.new("RGB", (1800, 930), "white")
    draw = ImageDraw.Draw(image)
    centered(draw, (900, 82), "当前页面中的题目", font(42, True), "#17365D")
    centered(draw, (900, 132), "在第一个能够证明答案合理的阶段退出", font(26), "#5F6B76")
    stages = [
        (70, "阶段 0", "精确层级标识符", "996 组对应关系", "#17365D"),
        (500, "阶段 1", "章节可单调对齐", "限定候选范围", "#2F6B73"),
        (930, "阶段 2", "运算符上下文", "内容具区分度", "#47798E"),
        (1360, "阶段 3", "有界单调对齐", "邻居提供位置支持", "#6B7280"),
    ]
    for idx, (x, title, line1, line2, color) in enumerate(stages):
        draw.rounded_rectangle((x, 280, x + 360, 560), radius=22, fill="#F5F7F9", outline=color, width=5)
        centered(draw, (x + 180, 342), title, font(34, True), color)
        centered(draw, (x + 180, 420), line1, font(25), "#20252B")
        centered(draw, (x + 180, 482), line2, font(23), "#5F6B76")
        if idx < 3:
            arrow(draw, (x + 365, 420), (x + 425, 420))
            centered(draw, (x + 395, 382), "未解决", font(18), "#7B8792")
    draw.rounded_rectangle((520, 700, 1280, 830), radius=20, fill="#FFF5F5", outline="#9B1C1C", width=4)
    centered(draw, (900, 765), "NONE：拒答，并返回候选项与原因", font(29, True), "#9B1C1C")
    for x, *_ in stages:
        draw.line([(x + 180, 565), (900, 698)], fill="#C1C7CD", width=2)
    image.save(out_dir / "figure_1.png", dpi=(220, 220))

    # Figure 2: precision and refusal rate.
    image = Image.new("RGB", (1700, 960), "white")
    draw = ImageDraw.Draw(image)
    left, right, top, bottom = 360, 1600, 190, 820
    centered(draw, (680, 70), "已接受匹配的精确率", font(26, True), "#17365D")
    draw.rectangle((420, 56, 460, 84), fill="#17365D")
    centered(draw, (1175, 70), "尝试拒答比例", font(26, True), "#6B4D00")
    draw.rectangle((930, 56, 970, 84), fill="#B08D57")
    for tick in range(0, 101, 25):
        x = left + int((right - left) * tick / 100)
        draw.line([(x, top), (x, bottom)], fill="#DDE2E6", width=2)
        centered(draw, (x, bottom + 44), f"{tick}%", font(21), "#5F6B76")
    regimes = ["两侧均有书签", "答案册无书签", "习题册无书签", "两侧均无书签"]
    precision = [100, 100, 100, 100]
    refusal = [8.9, 17.8, 98.0, 10.2]
    for idx, (label, p, r) in enumerate(zip(regimes, precision, refusal)):
        y = top + 80 + idx * 145
        draw.text((left - 25, y + 22), label, font=font(25), fill="#20252B", anchor="rm")
        x_p = left + int((right - left) * p / 100)
        draw.rectangle((left, y, x_p, y + 36), fill="#17365D")
        draw.text((x_p - 12, y + 18), f"{p}%", font=font(21, True), fill="white", anchor="rm")
        x_r = left + int((right - left) * r / 100)
        if r > 0:
            draw.rectangle((left, y + 50, x_r, y + 86), fill="#B08D57")
        draw.text((x_r + 12, y + 68), f"{r:g}%", font=font(21, True), fill="#6B4D00", anchor="lm")
    centered(draw, ((left + right) // 2, 912), "比例（%）", font(23), "#5F6B76")
    image.save(out_dir / "figure_2.png", dpi=(220, 220))

    # Figure 3: latency, read from the measurement rather than typed in here.
    #
    # This block used to hardcode [0, 342, 507, 45] against a 500 ms target, and
    # those numbers went stale silently the moment the retrieval path changed.
    # They now come from figures/latency-by-regime.data.json, which
    # tools/measure-regimes.mjs writes; if it is absent the build stops rather
    # than drawing a figure nobody can trace. See section 5.5 of the report.
    data_path = Path(__file__).resolve().parent / "figures" / "latency-by-regime.data.json"
    if not data_path.exists():
        raise SystemExit(
            """missing %s
Regenerate it with:
  node tools/measure-regimes.mjs <corpus> --json tmp/regimes.json
  node tools/make-figures.mjs --from tmp/regimes.json""" % data_path
        )
    latency = json.loads(data_path.read_text(encoding="utf-8"))
    zh_label = {
        "both": "两侧均有书签",
        "ansNone": "答案册无书签",
        "exNone": "习题册无书签",
        "neither": "两侧均无书签",
    }
    deadline = latency["deadlineMs"]
    target = latency["targetMs"]

    rows = latency["regimes"]
    observed_max = max(deadline, *(regime["maxP95"] for regime in rows))

    image = Image.new("RGB", (1700, 980), "white")
    draw = ImageDraw.Draw(image)
    left, right, top, bottom = 430, 1610, 150, 780
    xmax = max(1800, ((observed_max + 299) // 300) * 300)

    def px(value):
        return left + int((right - left) * value / xmax)

    for tick in range(0, xmax + 1, 300):
        x = px(tick)
        draw.line([(x, top), (x, bottom)], fill="#DDE2E6", width=2)
        draw.text((x, bottom + 40), str(tick), font=font(21), fill="#5F6B76", anchor="ma")

    # The alignment deadline. Reaching it means the result came from expiry
    # rather than from a decision, which is why it is drawn as a status line.
    xd = px(deadline)
    for y0 in range(top - 20, bottom, 22):
        draw.line([(xd, y0), (xd, min(y0 + 11, bottom))], fill="#9B1C1C", width=4)
    draw.text((xd - 12, top - 30), "%d ms 对齐截止时间" % deadline,
              font=font(23, True), fill="#9B1C1C", anchor="rb")

    xt = px(target)
    for y0 in range(top, bottom, 18):
        draw.line([(xt, y0), (xt, min(y0 + 9, bottom))], fill="#B08D57", width=3)
    draw.text((xt + 10, bottom + 76), "%d ms 响应目标" % target,
              font=font(21), fill="#6B4D00", anchor="la")

    step = (bottom - top) // max(len(rows), 1)
    for idx, regime in enumerate(rows):
        cy = top + step * idx + step // 2
        lo, hi = regime["minP95"], regime["maxP95"]
        x0, x1 = px(lo), px(hi)
        draw.rounded_rectangle((x0, cy - 16, max(x1, x0 + 6), cy + 16), radius=16, fill="#BBD3EC")
        for pair in regime["pairs"]:
            xp = px(pair["p95"])
            draw.ellipse((xp - 10, cy - 10, xp + 10, cy + 10), fill="#17365D")
        draw.text((left - 25, cy), zh_label.get(regime["key"], regime["key"]),
                  font=font(25), fill="#20252B", anchor="rm")
        text = "%d ms" % hi if lo == hi else "%d–%d ms" % (lo, hi)
        draw.text((max(x1, x0 + 6) + 16, cy), text, font=font(22, True), fill="#20252B", anchor="lm")

    draw.line([(left, top), (left, bottom)], fill="#9AA4AE", width=3)
    draw.text((left, top - 70), "页面级 p95 延迟（ms），3 组配对上的区间",
              font=font(23), fill="#5F6B76", anchor="ls")
    image.save(out_dir / "figure_3.png", dpi=(220, 220))

    # Figure 4: window-radius sweep.
    image = Image.new("RGB", (1700, 1020), "white")
    draw = ImageDraw.Draw(image)
    left, right, top, bottom = 190, 1590, 150, 820
    radii = [1, 2, 3, 4, 5, 6]
    mean = [0.291, 0.452, 0.652, 0.524, 0.540, 0.338]
    worst = [0.167, 0.250, 0.500, 0.333, 0.333, 0.000]
    ymax = 0.75
    for tick in [0, 0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7]:
        y = bottom - int((bottom - top) * tick / ymax)
        draw.line([(left, y), (right, y)], fill="#DDE2E6", width=2)
        draw.text((left - 20, y), f"{tick:.1f}", font=font(20), fill="#5F6B76", anchor="rm")
    points_mean, points_worst = [], []
    for i, radius in enumerate(radii):
        x = left + int((right - left) * i / 5)
        draw.text((x, bottom + 38), str(radius), font=font(22), fill="#20252B", anchor="ma")
        points_mean.append((x, bottom - int((bottom - top) * mean[i] / ymax)))
        points_worst.append((x, bottom - int((bottom - top) * worst[i] / ymax)))
    draw.line(points_mean, fill="#17365D", width=6, joint="curve")
    draw.line(points_worst, fill="#2F6B73", width=5, joint="curve")
    for point in points_mean:
        draw.ellipse((point[0] - 9, point[1] - 9, point[0] + 9, point[1] + 9), fill="#17365D")
    for point in points_worst:
        draw.rectangle((point[0] - 8, point[1] - 8, point[0] + 8, point[1] + 8), fill="#2F6B73")
    for value, color, label in [(0.145, "#7B8792", "普通相似度 0.145"), (0.184, "#B08D57", "数学片段二元组 0.184")]:
        y = bottom - int((bottom - top) * value / ymax)
        for x in range(left, right, 28):
            draw.line([(x, y), (min(x + 14, right), y)], fill=color, width=3)
        draw.text((right - 8, y - 8), label, font=font(19), fill=color, anchor="rb")
    chosen = points_mean[2]
    draw.ellipse((chosen[0] - 15, chosen[1] - 15, chosen[0] + 15, chosen[1] + 15), outline="#B08D57", width=7)
    draw.text((chosen[0] + 25, chosen[1] - 30), "选定半径", font=font(22, True), fill="#6B4D00", anchor="ls")
    inverted = points_worst[5]
    draw.line([(inverted[0] - 14, inverted[1] - 14), (inverted[0] + 14, inverted[1] + 14)], fill="#9B1C1C", width=6)
    draw.line([(inverted[0] + 14, inverted[1] - 14), (inverted[0] - 14, inverted[1] + 14)], fill="#9B1C1C", width=6)
    draw.text((inverted[0] - 10, inverted[1] - 25), "1 对样本反转", font=font(22, True), fill="#9B1C1C", anchor="rb")
    draw.rectangle((300, 54, 338, 80), fill="#17365D")
    draw.text((350, 68), "平均间隔", font=font(22), fill="#20252B", anchor="lm")
    draw.rectangle((650, 54, 688, 80), fill="#2F6B73")
    draw.text((700, 68), "最差情况", font=font(22), fill="#20252B", anchor="lm")
    centered(draw, ((left + right) // 2, 930), "运算符上下文窗口半径（每侧字符数）", font(23), "#5F6B76")
    draw.text((left, top - 30), "间隔（正确候选 - 最佳错误候选）", font=font(22), fill="#5F6B76", anchor="ls")
    image.save(out_dir / "figure_4.png", dpi=(220, 220))


def add_figure(doc: Document, path: Path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        parts = [x.strip() for x in lines[i].strip().strip("|").split("|")]
        if i != start + 1:
            rows.append(parts)
        i += 1
    return rows, i


def build(markdown_path: Path, output_path: Path, figures_dir: Path):
    make_figures(figures_dir)
    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    configure_section(section)
    configure_styles(doc)
    add_running_furniture(section)
    add_cover(doc, *read_version(markdown_path))
    num_id = add_numbering_definition(doc)

    doc.core_properties.title = "Find-Engine：面向成对数学 PDF 的确定性题目-答案对齐"
    doc.core_properties.subject = "确定性题目-答案对齐技术研究报告"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "确定性对齐; PDF 文档结构; 题目答案匹配; 数学相似度; 拒答机制"
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    start = next(i for i, line in enumerate(lines) if line.startswith("## 摘要"))
    i = start
    table_no = 0
    in_abstract = False

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line or line == "---":
            i += 1
            continue

        if line.startswith("## "):
            title = line[3:].strip()
            in_abstract = title == "摘要"
            p = doc.add_paragraph(style="Heading 1")
            if in_abstract:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(10)
            elif title.startswith("附录 A"):
                p.paragraph_format.page_break_before = True
            add_inline(p, title, size=15, color=NAVY, cjk=HEAD_CJK)
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip(), size=12.5, color=TEAL, cjk=HEAD_CJK)
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        figure_match = FIGURE_RE.fullmatch(line)
        if figure_match:
            figure_no = int(figure_match.group(1) or figure_match.group(2))
            add_figure(doc, figures_dir / f"figure_{figure_no}.png")
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows, i = parse_table(lines, i)
            table_no += 1
            add_table(doc, rows, table_no)
            continue

        if line.startswith("```"):
            language = line[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.right_indent = Cm(0.35)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.08
            p.paragraph_format.keep_together = True
            shade_paragraph(p, "F3F5F6")
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, cjk=BODY_CJK, latin=MONO, size=8.8, color=INK)
            continue

        if re.match(r"^\d+\.\s+", line):
            content = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph()
            apply_numbering(p, num_id)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.35
            add_inline(p, content)
            i += 1
            continue

        # Gather a normal paragraph until the next structural marker.
        parts = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if (
                nxt == "---"
                or nxt.startswith("## ")
                or nxt.startswith("### ")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or FIGURE_RE.fullmatch(nxt)
                or re.match(r"^\d+\.\s+", nxt)
            ):
                break
            parts.append(nxt)
            i += 1
        content = " ".join(parts)

        is_fig_caption = bool(re.match(r"^\*\*图\s*\d+", content))
        is_keywords = content.startswith("**关键词：**")
        is_figure_index = content.startswith("**插图索引：**")
        is_note = content.startswith("†")
        is_labelled = bool(re.match(r"^\*\*[^*]+。?\*\*", content))

        if is_fig_caption:
            p = doc.add_paragraph(style="Caption")
            p.paragraph_format.keep_together = True
            add_inline(p, content, size=9, color=MUTED)
        elif is_keywords:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, content, size=9.5, color=TEAL, cjk=HEAD_CJK)
        elif is_figure_index:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.right_indent = Cm(0.35)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(9)
            shade_paragraph(p, LIGHT_BLUE)
            add_inline(p, content, size=9.3, color=NAVY)
        elif is_note:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, content, size=8.8, color=MUTED)
        else:
            style = "Abstract Body" if in_abstract else "Normal"
            p = doc.add_paragraph(style=style)
            if is_labelled:
                p.paragraph_format.keep_together = True
            add_inline(p, content)

    # Keep source code examples and tables intact while avoiding accidental track metadata.
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--figures-dir", type=Path, required=True)
    args = parser.parse_args()
    build(args.markdown, args.output, args.figures_dir)


if __name__ == "__main__":
    main()
